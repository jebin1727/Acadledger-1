
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os

source_path = r"d:\Final-Year-Project\docs\Project Report Format.docx"
dest_path = r"d:\Final-Year-Project\docs\Project_Report_AcadLedger.docx"

if not os.path.exists(source_path):
    print("Source file not found.")
    exit()

doc = docx.Document(source_path)

# --- Content Definitions ---

CONTENT = {
    "ABSTRACT": """Academic credential fraud, including the fabrication of degree certificates and grade sheets, poses a significant trust deficit. Traditional verification is manual and slow. 
    
**AcadLedger** is a Privacy-Preserving Attestation Framework. It uses **Generative AI (Gemini 1.5)** to detect visual tampering (Layer 1) and **Polygon Blockchain** to anchor semantic data hashes (Layer 2). This "Double Shield" ensures that valid credentials can be verified even if file formats change (PDF vs Scan), while preventing data from being stored on public servers.""",

    "INTRODUCTION": """### 1.1 Overview
"AcadLedger" is a dApp designed to issue and verify academic credentials using AI and Blockchain. It moves the trust anchor to a transparent ledger while using AI to validatethe visual integrity of documents.

### 1.2 Problem Statement
1. **Fake Degrees**: High-quality forgeries are hard to detect.
2. **Privacy**: Storing PII on public blockchains violates GDPR.
3. **File Hashing Limit**: Hashing files fails if the file is compressed or scanned.
4. **Centralization**: University servers are single points of failure.

### 1.3 Objectives
- **Privacy**: No PII on-chain.
- **Double Shield**: Visual checking (AI) + Factual checking (Blockchain).
- **Robustness**: Verify data, not just files.
- **Scalability**: Low cost via Polygon.

### 1.4 Scope
Focuses on Academic Degrees but extensible to Medical Records and IDs.""",

    "LITERATURESURVEY": """### 2.1 Traditional Systems
Manual verification is slow (weeks) and expensive.
### 2.2 Public Blockchains
Storing data on Ethereum is expensive and not private.
### 2.3 Soulbound Tokens
SBTs lack a bridge for physical paper credentials.
### 2.4 Existing File Hashing
DocuSign-like solutions fail if you print and scan the document (hash mismatch).
**AcadLedger Gap**: We use Semantic Hashing to fix the robustness issue.""",

    "SYSTEMANALYSIS": """### 3.1 Existing System
Employers wait weeks for verification. Students lose control of data.
### 3.2 Proposed System
Instant verification. Student owns data. AI detects tampering.
### 3.3 Feasibility
Technical: Gemini + Polygon are production-ready.
Economic: Cheaper than maintaining secure servers.""",

    "SYSTEMDESIGN": """### 4.1 Architecture
**Data (Off-Chain)**: Student holds the PDF.
**Proof (On-Chain)**: Blockchain holds the Hash.
**Bridge**: AI extracts data and hashes it.

### 4.2 Modules
1. **Attestation**: Minting credentials after AI check.
2. **Verification**: Checking credentials against the chain.
3. **AI Guard**: Detecting visual fraud.

### 4.3 Algorithms
**Semantic Hashing**:
1. AI Extracts Data -> JSON.
2. Sort JSON Keys.
3. Hash(JSON) -> Store on Chain.""",

    "SYSTEMIMPLEMENTATION": """### 5.1 Tech Stack
- **Frontend**: Next.js, Tailwind.
- **AI**: Google Gemini.
- **Chain**: Polygon Amoy.

### 5.2 Implementation
We built a web client that allows drag-and-drop verification. The "AI Guard" analyzes the pixels before hashing, ensuring no "photoshopped" degrees are anchored.""",

    "CONCLUSION": """### 6.1 Conclusion
AcadLedger successfully decouples Data from Proof, ensuring privacy and robustness.
### 6.2 Future Work
Mobile App with AR scanning and ZK-Proofs for granular privacy."""
}

# --- Logic ---

def add_markdown_content(doc, header_paragraph, content_text):
    """
    Inserts formatted content after the given header paragraph.
    """
    current_p = header_paragraph
    
    # Simple split by lines
    lines = content_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        new_p = doc.add_paragraph()
        
        # Heading handling
        if line.startswith("###"):
            text = line.replace("###", "").strip()
            run = new_p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            # Fallback: Just use text with a bullet character if style is missing
            try:
                new_p.style = 'List Bullet'
            except:
                pass # Styles might not exist
            apply_formatting(new_p, text)
        elif line[0].isdigit() and ". " in line[:4]:
            text = line # Keep the number
            try:
                new_p.style = 'List Number'
                text = line.split(". ", 1)[1].strip() # Remove number if style handles it
            except:
                pass # Use full text if style fails
            apply_formatting(new_p, text)
        else:
            # Normal text
            apply_formatting(new_p, line)
            new_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Move new_p after current_p
        current_p._element.addnext(new_p._element)
        current_p = new_p

def apply_formatting(paragraph, text):
    """
    Parses **bold** markdown.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)

def replace_section_content(doc, trigger_keywords, content_key):
    """
    Finds a paragraph matching ANY of the trigger keywords.
    Deletes subsequent content until the next major Header.
    Inserts new content.
    """
    start_paragraph = None
    start_index = -1
    
    # 1. Find Start
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper().replace(" ", "")
        for kw in trigger_keywords:
            if kw.upper().replace(" ", "") in text:
                start_paragraph = p
                start_index = i
                print(f"Found Section: {kw} at index {i}")
                break
        if start_paragraph:
            break
            
    if not start_paragraph:
        print(f"Could not find section for {content_key}")
        return

    # 2. Find End (Next Chapter or significant header)
    end_index = len(doc.paragraphs)
    for i in range(start_index + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip().upper()
        # Heuristic for next section
        if "CHAPTER" in text or "TABLE OF" in text or "REFERENCES" in text:
            end_index = i
            break
        # Also check for specific known headers if "Chapter" isn't used consistently
        if "INTRODUCTION" in text and content_key == "ABSTRACT": end_index = i; break
        if "LITERATURE" in text and content_key == "INTRODUCTION": end_index = i; break
        if "SYSTEM" in text and content_key == "LITERATURESURVEY": end_index = i; break

    # 3. Delete old content
    # We collect paragraphs to delete to avoid modifying list while iterating
    # Note: doc.paragraphs is a live list somewhat, but accessing by index from original list object is safer if we don't re-query.
    # Actually, removing elements from the XML tree is the robust way.
    
    # We will identify paragraphs to remove
    paragraphs_to_remove = []
    # We grab them from the document's internal list
    for i in range(start_index + 1, end_index):
        paragraphs_to_remove.append(doc.paragraphs[i])
        
    for p in paragraphs_to_remove:
        p_element = p._element
        if p_element.getparent() is not None:
            p_element.getparent().remove(p_element)

    # 4. Insert New Content
    add_markdown_content(doc, start_paragraph, CONTENT[content_key])


# --- Execution Order ---
# We map keys to likely headers in the doc
replace_section_content(doc, ["ABSTRACT"], "ABSTRACT")
replace_section_content(doc, ["INTRODUCTION", "CHAPTER 1"], "INTRODUCTION")
replace_section_content(doc, ["LITERATURE", "CHAPTER 2"], "LITERATURESURVEY")
replace_section_content(doc, ["SYSTEM ANALYSIS", "CHAPTER 3"], "SYSTEMANALYSIS") # Assumption on Chapter 3
replace_section_content(doc, ["SYSTEM DESIGN", "CHAPTER 4"], "SYSTEMDESIGN") # Assumption
replace_section_content(doc, ["IMPLEMENTATION", "CHAPTER 5"], "SYSTEMIMPLEMENTATION") # Assumption
replace_section_content(doc, ["CONCLUSION", "CHAPTER 6"], "CONCLUSION")

doc.save(dest_path)
print(f"Done. Saved to {dest_path}")
